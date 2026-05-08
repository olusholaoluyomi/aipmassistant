#!/usr/bin/env python3
"""
market_intel_focused.py — Focused 7-day digest: AI Marketing + AI Care only.
Output: intel/YYYY-MM-DD_HHMM.docx
"""

import os, re, sys
from datetime import datetime
from pathlib import Path

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

import anthropic
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MODEL = "claude-opus-4-6"
DAYS  = 7

COLOR_GREEN = RGBColor(0x1A, 0xD6, 0x78)
COLOR_TEAL  = RGBColor(0x00, 0x9E, 0xA6)
COLOR_DARK  = RGBColor(0x0F, 0x17, 0x2A)
COLOR_MUTED = RGBColor(0x94, 0xA3, 0xB8)

LAYERS = [
    {
        "id": "agentic_marketing",
        "title": "Agentic Marketing & Campaign Orchestration",
        "competitors": ["Braze", "MoEngage", "WebEngage", "CleverTap", "Insider"],
        "focus": (
            "AI-driven campaign features, autonomous audience segmentation, agentic workflows, "
            "generative AI for content, RFM/predictive models, and marketing automation updates."
        ),
    },
    {
        "id": "ai_customer_care",
        "title": "AI Customer Care & CCaaS",
        "competitors": ["Intercom", "Zendesk", "Genesys", "Salesforce Service Cloud", "Freshworks"],
        "focus": (
            "AI agent/copilot features, voice AI, contact center automation, LLM integrations, "
            "CCaaS platform updates, agent assist, and conversational AI launches."
        ),
    },
]


def build_system(run_date: str) -> str:
    return (
        f"You are a senior competitive intelligence analyst for Unifonic, a leading AI-native CX platform in Saudi Arabia and the GCC.\n"
        f"Unifonic's core products: MCC (marketing campaigns), Flow Studio (journey automation), Agent Console (live chat), "
        f"Chatbot Builder, WhatsApp/SMS/Voice APIs, and an Agentic Marketing framework with AI models.\n"
        f"Today's date is {run_date}. Be concise, factual, and commercially sharp. "
        f"Include source URLs inline where available. Focus on what is actionable for Unifonic's product and GTM strategy."
    )


def build_prompt(layer: dict, run_date: str) -> str:
    competitors_str = ", ".join(layer["competitors"])
    return (
        f"Today is {run_date}. Search the web for competitive intelligence from the past {DAYS} days on: {competitors_str}.\n\n"
        f"Focus area: {layer['focus']}\n\n"
        f"Sources to prioritize: company blogs, official changelogs, press releases, G2 reviews, TechCrunch, VentureBeat, "
        f"LinkedIn announcements, Gartner peer insights, and industry newsletters.\n\n"
        f"For each competitor, return exactly this format:\n"
        f"### **[Competitor Name]**\n"
        f"- **What's new:** [1-2 bullets — specific update with date. Include source URL in parentheses where available.]\n"
        f"- **Strategic signal:** [1 bullet — what this means for Unifonic]\n\n"
        f"If no recent news in the past {DAYS} days, write 'No significant updates in the past {DAYS} days.'\n"
        f"Do NOT add introductory text, date headers, or section titles. Start directly with the first competitor.\n"
        f"Keep total response under 600 words."
    )


def _extract_text(response) -> str:
    return "\n".join(b.text for b in response.content if hasattr(b, "text"))


def run_search(client, prompt: str, system: str, max_tokens: int = 2048) -> str:
    messages = [{"role": "user", "content": prompt}]
    for _ in range(5):
        response = client.messages.create(
            model=MODEL,
            max_tokens=max_tokens,
            thinking={"type": "adaptive"},
            system=system,
            tools=[{"type": "web_search_20260209", "name": "web_search"}],
            messages=messages,
        )
        if response.stop_reason == "end_turn":
            return _extract_text(response)
        if response.stop_reason == "pause_turn":
            messages = [
                {"role": "user", "content": prompt},
                {"role": "assistant", "content": response.content},
            ]
            continue
        return _extract_text(response) or f"[No output — stop_reason: {response.stop_reason}]"
    return "[Search loop exceeded max continuations]"


def _add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_paragraph_with_links(doc, text: str, style: str = "Normal", bold_prefix: str = None):
    url_pattern = re.compile(r'(https?://[^\s\)\]]+)')
    para = doc.add_paragraph(style=style)
    remaining = text
    if bold_prefix and text.startswith(bold_prefix):
        run = para.add_run(bold_prefix)
        run.bold = True
        run.font.color.rgb = COLOR_DARK
        remaining = text[len(bold_prefix):]
    for part in url_pattern.split(remaining):
        if url_pattern.match(part):
            _add_hyperlink(para, part, part)
        elif part:
            run = para.add_run(part)
            run.font.color.rgb = COLOR_DARK
    return para


def build_docx(layer_results, run_ts: str) -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    title = doc.add_heading("Unifonic Competitive Intelligence — AI Focus", level=0)
    title.runs[0].font.color.rgb = COLOR_TEAL

    sub = doc.add_paragraph(f"Scope: AI Marketing & AI Care  |  Last {DAYS} days  |  Generated: {run_ts}")
    sub.runs[0].font.color.rgb = COLOR_MUTED
    sub.runs[0].font.italic = True
    doc.add_paragraph()

    for _layer_id, layer_title, content in layer_results:
        h2 = doc.add_heading(layer_title, level=1)
        h2.runs[0].font.color.rgb = COLOR_TEAL

        for line in content.strip().splitlines():
            line = line.strip()
            if not line:
                continue
            if line.startswith("###"):
                name = re.sub(r"[#*`]", "", line).strip()
                h3 = doc.add_heading(name, level=2)
                h3.runs[0].font.color.rgb = COLOR_GREEN
            elif line.startswith("- "):
                bullet_text = line[2:].strip()
                bold_match = re.match(r"\*\*(.+?)\*\*\s*(.*)", bullet_text, re.DOTALL)
                if bold_match:
                    prefix = bold_match.group(1) + " "
                    rest = bold_match.group(2)
                    _add_paragraph_with_links(doc, prefix + rest, style="List Bullet", bold_prefix=prefix)
                else:
                    _add_paragraph_with_links(doc, bullet_text, style="List Bullet")
            elif line.startswith("  - "):
                _add_paragraph_with_links(doc, line[4:].strip(), style="List Bullet 2")
            else:
                clean = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
                _add_paragraph_with_links(doc, clean)

        doc.add_paragraph()

    note = doc.add_paragraph("Generated by market_intel_focused.py using Anthropic Web Search.")
    note.runs[0].font.color.rgb = COLOR_MUTED
    note.runs[0].font.italic = True
    note.runs[0].font.size = Pt(9)

    return doc


def main():
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        print("ERROR: ANTHROPIC_API_KEY not set.", file=sys.stderr)
        sys.exit(1)

    client = anthropic.Anthropic(api_key=api_key)
    intel_dir = Path(__file__).parent / "intel"
    intel_dir.mkdir(exist_ok=True)

    run_ts = datetime.now().strftime("%Y-%m-%d %H:%M")
    file_ts = datetime.now().strftime("%Y-%m-%d_%H%M")
    output_path = intel_dir / f"{file_ts}_ai_focus.docx"

    print(f"Unifonic AI Market Intel — {run_ts}  (last {DAYS} days)")
    print(f"Output: {output_path}\n")

    system = build_system(run_ts)
    layer_results = []

    for layer in LAYERS:
        print(f"[{layer['id']}] Searching: {layer['title']} ...")
        try:
            content = run_search(client, build_prompt(layer, run_ts), system)
            layer_results.append((layer["id"], layer["title"], content))
            print("  Done.")
        except anthropic.APIError as e:
            layer_results.append((layer["id"], layer["title"], f"[Search failed: {e}]"))
            print(f"  ERROR: {e}", file=sys.stderr)

    doc = build_docx(layer_results, run_ts)
    doc.save(output_path)
    print(f"\nDigest saved to: {output_path}")


if __name__ == "__main__":
    main()
