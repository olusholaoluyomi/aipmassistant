#!/usr/bin/env python3
"""
market_intel_digest.py — Unifonic Monthly Competitive Intelligence
Generates a Word (.docx) digest of competitor activity across 4 layers.
Output: intel/YYYY-MM-DD_HHMM.docx
"""

import os
import re
import sys
from datetime import datetime
from pathlib import Path

# Load .env if present (pip install python-dotenv)
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

import anthropic
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------

MODEL = "claude-opus-4-6"

# Unifonic brand colors
COLOR_GREEN = RGBColor(0x1A, 0xD6, 0x78)  # #1AD678 — Primary accent
COLOR_TEAL  = RGBColor(0x00, 0x9E, 0xA6)  # #009EA6 — Secondary accent (headings)
COLOR_DARK  = RGBColor(0x0F, 0x17, 0x2A)  # #0F172A — Primary text
COLOR_MUTED = RGBColor(0x94, 0xA3, 0xB8)  # #94A3B8 — Neutral / muted text

LAYERS = [
    {
        "id": "cpaas",
        "title": "CPaaS & Channels",
        "competitors": ["Twilio", "Infobip", "Sinch", "MessageBird", "T2 (STC)"],
        "focus": (
            "New channel integrations, API releases, pricing changes, platform updates, "
            "partnership announcements, and developer ecosystem moves in CPaaS and messaging."
        ),
    },
    {
        "id": "agentic_marketing",
        "title": "Agentic Marketing & Campaign Orchestration",
        "competitors": ["Braze", "MoEngage", "WebEngage", "CleverTap", "Insider"],
        "focus": (
            "AI-driven campaign features, autonomous audience segmentation, agentic workflows, "
            "generative AI for content, RFM/predictive models, and marketing automation updates."
        ),
    },
    {
        "id": "ai_customer_care",
        "title": "AI Customer Care & CCaaS",
        "competitors": ["Intercom", "Zendesk", "Genesys", "Salesforce Service Cloud", "Freshworks"],
        "focus": (
            "AI agent/copilot features, voice AI, contact center automation, LLM integrations, "
            "CCaaS platform updates, agent assist, and conversational AI launches."
        ),
    },
    {
        "id": "mena_regional",
        "title": "MENA & Regional Players",
        "competitors": ["Wati", "STC", "Zain", "Mobily", "Taqnyat"],
        "focus": (
            "Regional product launches, Saudi/GCC market moves, telco digital offerings, "
            "WhatsApp BSP developments, and regional AI/CX announcements."
        ),
    },
]

MENA_SIGNALS_PROMPT = """
Search for recent news and signals relevant to the GCC/MENA technology and CPaaS market:
- AI CX and conversational AI adoption in Saudi Arabia and UAE
- WhatsApp Business API ecosystem news in GCC
- Regulatory updates (CITC, TRA) affecting messaging or CPaaS
- Enterprise digital transformation deals in Saudi Arabia or UAE
- Gartner/Forrester analyst notes on CPaaS, CCaaS, or AI CX platforms
- Funding rounds or M&A in MENA tech (CPaaS, AI, CX)
- Key events: GITEX, LEAP, Saudi Vision 2030 digital initiatives

Return a concise section with 3-5 bullet points covering the most strategically relevant signals for Unifonic.
Each bullet: signal headline + why it matters to Unifonic. Include source URLs where available.
"""


# ---------------------------------------------------------------------------
# Prompts
# ---------------------------------------------------------------------------

def build_layer_system_prompt(run_date: str) -> str:
    return (
        f"You are a senior competitive intelligence analyst for Unifonic, a leading AI-native CX platform in Saudi Arabia and the GCC.\n"
        f"Unifonic's core products: MCC (marketing campaigns), Flow Studio (journey automation), Agent Console (live chat), "
        f"Chatbot Builder, WhatsApp/SMS/Voice APIs, and an Agentic Marketing framework with AI models.\n"
        f"Today's date is {run_date}. Do not add any date headers or section titles — the digest template handles that. "
        f"Be concise, factual, and commercially sharp. Include source URLs inline where available. "
        f"Focus on what is actionable for Unifonic's product and GTM strategy."
    )


def build_layer_prompt(layer: dict, run_date: str) -> str:
    competitors_str = ", ".join(layer["competitors"])
    return (
        f"Today is {run_date}. Search the web for competitive intelligence from the past 30 days on: {competitors_str}.\n\n"
        f"Focus area: {layer['focus']}\n\n"
        f"Sources to prioritize: company blogs, official changelogs, press releases, G2 reviews, TechCrunch, VentureBeat, "
        f"LinkedIn announcements, Gartner peer insights, and industry newsletters.\n\n"
        f"For each competitor, return exactly this format:\n"
        f"### **[Competitor Name]**\n"
        f"- **What's new:** [1-2 bullets — specific update with date. Include source URL in parentheses where available.]\n"
        f"- **Strategic signal:** [1 bullet — what this means for Unifonic]\n\n"
        f"If no recent news, write 'No significant updates in the past 30 days.'\n"
        f"Do NOT add introductory text, date headers, or section titles. Start directly with the first competitor.\n"
        f"Keep total response under 600 words."
    )


# ---------------------------------------------------------------------------
# API calls
# ---------------------------------------------------------------------------

def _extract_text(response) -> str:
    return "\n".join(b.text for b in response.content if hasattr(b, "text"))


def run_search(client: anthropic.Anthropic, prompt: str, system: str, max_tokens: int = 2048) -> str:
    messages = [{"role": "user", "content": prompt}]
    max_continuations = 5

    for _ in range(max_continuations):
        response = client.messages.create(
            model=MODEL,
            max_tokens=max_tokens,
            thinking={"type": "adaptive"},
            system=system,
            tools=[{"type": "web_search_20260209", "name": "web_search"}],
            messages=messages,
        )
        if response.stop_reason == "end_turn":
            return _extract_text(response)
        if response.stop_reason == "pause_turn":
            messages = [
                {"role": "user", "content": prompt},
                {"role": "assistant", "content": response.content},
            ]
            continue
        return _extract_text(response) or f"[No output — stop_reason: {response.stop_reason}]"

    return "[Search loop exceeded max continuations]"


# ---------------------------------------------------------------------------
# Word document builder
# ---------------------------------------------------------------------------

def _add_hyperlink(paragraph, text: str, url: str):
    """Add a clickable hyperlink run to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_paragraph_with_links(doc: Document, text: str, style: str = "Normal", bold_prefix: str = None):
    """Add a paragraph, converting any URLs in the text to clickable hyperlinks."""
    url_pattern = re.compile(r'(https?://[^\s\)\]]+)')
    para = doc.add_paragraph(style=style)

    # Handle bold prefix like "What's new:" or "Strategic signal:"
    remaining = text
    if bold_prefix and text.startswith(bold_prefix):
        run = para.add_run(bold_prefix)
        run.bold = True
        run.font.color.rgb = COLOR_DARK
        remaining = text[len(bold_prefix):]

    parts = url_pattern.split(remaining)
    for part in parts:
        if url_pattern.match(part):
            _add_hyperlink(para, part, part)
        else:
            if part:
                run = para.add_run(part)
                run.font.color.rgb = COLOR_DARK
    return para


def build_docx(layer_results: list, mena_signals: str, run_ts: str) -> Document:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Title
    title = doc.add_heading("Unifonic Competitive Intelligence Digest", level=0)
    title.runs[0].font.color.rgb = COLOR_TEAL

    # Subtitle / date
    sub = doc.add_paragraph(f"Generated: {run_ts}")
    sub.runs[0].font.color.rgb = COLOR_MUTED
    sub.runs[0].font.italic = True

    doc.add_paragraph()  # spacer

    # Each layer
    for _layer_id, layer_title, content in layer_results:
        h2 = doc.add_heading(layer_title, level=1)
        h2.runs[0].font.color.rgb = COLOR_TEAL

        for line in content.strip().splitlines():
            line = line.strip()
            if not line:
                continue

            # Competitor heading: ### **Name** or ### Name
            if line.startswith("###"):
                name = re.sub(r"[#*`]", "", line).strip()
                h3 = doc.add_heading(name, level=2)
                h3.runs[0].font.color.rgb = COLOR_GREEN

            # Bullet point
            elif line.startswith("- "):
                bullet_text = line[2:].strip()
                # Detect bold prefix like **What's new:**
                bold_match = re.match(r"\*\*(.+?)\*\*\s*(.*)", bullet_text, re.DOTALL)
                if bold_match:
                    prefix = bold_match.group(1) + " "
                    rest = bold_match.group(2)
                    _add_paragraph_with_links(doc, prefix + rest, style="List Bullet", bold_prefix=prefix)
                else:
                    _add_paragraph_with_links(doc, bullet_text, style="List Bullet")

            # Sub-bullet
            elif line.startswith("  - "):
                _add_paragraph_with_links(doc, line[4:].strip(), style="List Bullet 2")

            # Plain text
            else:
                clean = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
                _add_paragraph_with_links(doc, clean)

        doc.add_paragraph()  # spacer between layers

    # MENA & Industry Signals
    h2 = doc.add_heading("MENA & Industry Signals", level=1)
    h2.runs[0].font.color.rgb = COLOR_TEAL

    for line in mena_signals.strip().splitlines():
        line = line.strip()
        if not line:
            continue
        if line.startswith("###"):
            name = re.sub(r"[#*`]", "", line).strip()
            h3 = doc.add_heading(name, level=2)
            h3.runs[0].font.color.rgb = COLOR_GREEN
        elif line.startswith("- ") or line.startswith("* "):
            _add_paragraph_with_links(doc, line[2:].strip(), style="List Bullet")
        elif line.startswith("#"):
            clean = re.sub(r"^#+\s*", "", line)
            doc.add_heading(clean, level=3)
        else:
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
            _add_paragraph_with_links(doc, clean)

    # Footer note
    doc.add_paragraph()
    note = doc.add_paragraph("Generated by market_intel_digest.py using Anthropic Web Search.")
    note.runs[0].font.color.rgb = COLOR_MUTED
    note.runs[0].font.italic = True
    note.runs[0].font.size = Pt(9)

    return doc


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        print("ERROR: ANTHROPIC_API_KEY environment variable not set.", file=sys.stderr)
        sys.exit(1)

    client = anthropic.Anthropic(api_key=api_key)

    intel_dir = Path(__file__).parent / "intel"
    intel_dir.mkdir(exist_ok=True)

    run_ts = datetime.now().strftime("%Y-%m-%d %H:%M")
    file_ts = datetime.now().strftime("%Y-%m-%d_%H%M")
    output_path = intel_dir / f"{file_ts}.docx"

    print(f"Unifonic Market Intel Digest — {run_ts}")
    print(f"Output: {output_path}")
    print()

    system = build_layer_system_prompt(run_ts)
    layer_results = []

    for layer in LAYERS:
        print(f"[{layer['id']}] Searching: {layer['title']} ...")
        try:
            content = run_search(client, build_layer_prompt(layer, run_ts), system)
            layer_results.append((layer["id"], layer["title"], content))
            print("  Done.")
        except anthropic.APIError as e:
            layer_results.append((layer["id"], layer["title"], f"[Search failed: {e}]"))
            print(f"  ERROR: {e}", file=sys.stderr)

    print("[mena] Searching: MENA & Industry Signals ...")
    try:
        mena_prompt = f"Today is {run_ts}.\n{MENA_SIGNALS_PROMPT}"
        mena_signals = run_search(client, mena_prompt, system, max_tokens=1024)
        print("  Done.")
    except anthropic.APIError as e:
        mena_signals = f"[Search failed: {e}]"
        print(f"  ERROR: {e}", file=sys.stderr)

    doc = build_docx(layer_results, mena_signals, run_ts)
    doc.save(output_path)

    print()
    print(f"Digest saved to: {output_path}")


if __name__ == "__main__":
    main()
